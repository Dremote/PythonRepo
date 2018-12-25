#ConvSPD.py, read from docx and fill in SpdTable
import docx
import getopt
import sys
import os
import xlrd
from win32com import client as wc

def _Crc16(spdbytes, count):
	crc = 0
	index = 0
	count -= 1;
	while count >= 0:
		crc = crc ^ (eval(spdbytes[index]) << 8)
		index += 1
		for i in range(8):
			if (crc & 0x8000):
				crc = crc << 1 ^ 0x1021
			else:
				crc = crc << 1
		count -= 1
	return (crc & 0xFFFF)
	
def _Crc16Check(SpdTable):
	data16 = _Crc16(SpdTable[0:128], 126)
	SPD_byte_126 = hex(data16 & 0xFF)
	SPD_byte_127 = hex(data16 >> 8)
	if not(eval(SPD_byte_126) == eval(SpdTable[126]) and eval(SPD_byte_127) == eval(SpdTable[127])):
		print("ERROR:SPD_byte_126, SPD_byte_127 CRC check fail!")
		return 1
	data16 = _Crc16(SpdTable[128:256], 126)
	SPD_byte_254 = hex(data16 & 0xFF)
	SPD_byte_255 = hex(data16 >> 8)
	if not(eval(SPD_byte_254) == eval(SpdTable[254]) and eval(SPD_byte_255) == eval(SpdTable[255])):
		print("ERROR:SPD_byte_254, SPD_byte_255 CRC check fail!")
		return 1
	print("CRC16 check pass.")
	return 0

class Micron:
	def __init__(self, DocxFile):
		self._DocxFile = DocxFile
		self._SpdTable = []
	
	def __GetSpdTableFromFile(self):
		CurDir = os.getcwd() + "\\"
		InputFile = CurDir + self._DocxFile
		
		#initialize SpdTable
		SpdTable = []
		for i in range(512):
			SpdTable.append(0)
		try:
			#Convert to format "DOCX"
			if InputFile[-3:].upper() == "DOC" or InputFile[-3:].upper() == "RTF":
				Word = wc.Dispatch('Word.Application')
				Doc = Word.Documents.Open(InputFile)
				Doc.SaveAs(InputFile[:-3] + "docx",16)
				Doc.Close()
				Word.Quit()
				DocxFile = docx.Document(InputFile[:-3] + "docx")
			else:
				DocxFile = docx.Document(InputFile)
		except:
			print("Please provide the file with valid format!")
			return 1
			
		#delete invalid lines
		FileRmvEmp = []
		for para in DocxFile.paragraphs:
			if para.text != "":
				FileRmvEmp.append(para.text)
				
		#fill in SpdTable	
		flag = 0;
		for lines in FileRmvEmp:
			if flag:
				bytes = lines.split()[0]
				values = lines.split()[-1]
				byte = bytes.split("-")
				if len(byte) == 1:
					index = eval(byte[0])
					SpdTable[index] = "0x" + values
				else:
					rangeleft = eval(byte[0])
					rangeright = eval(byte[1])
					if rangeleft == 329:
						#Manufacturer's part number written in ASCII format
						zeronum = (rangeright - rangeleft + 1) - len(values)
						for i in range(zeronum):
							values += " "
						for i in range(rangeleft, rangeright + 1):
							SpdTable[i] = "'" + values[i - rangeleft] + "'"
					else:
						if values.isnumeric() and eval(values) == 0:
							for i in range (rangeleft, rangeright + 1):
								SpdTable[i] = "0x00"
						else:
							for i in range(rangeleft, rangeright + 1):
								valuesindex = (i - rangeleft)*2
								SpdTable[i] = "0x" + values[valuesindex:valuesindex + 2]
			if lines[:4] == "BYTE":
				flag = 1;
				
		#CRC check
		ReturnVal = _Crc16Check(SpdTable)
		if not ReturnVal == 0:
			return ReturnVal
			
		self._SpdTable = SpdTable
		return 0
		
	def SaveSpdTableToFile(self, OutputFile):
		ReturnVal = self.__GetSpdTableFromFile()
		if not ReturnVal == 0:
			return ReturnVal
			
		#dump SpdTable
		SpdTable = self._SpdTable
		OutputFile = open(OutputFile, 'w')
		OutputFile.write("{//MICRON\n")
		OutputFile.write("//  ")
		for i in range(16):
			OutputFile.write("{:^5d}".format(i))
		OutputFile.write("\n")
		for i in range(len(SpdTable)):
			if (i % 16) == 0:
				OutputFile.write("    ")
				OutputFile.write("{:>4s}".format(SpdTable[i]) + ",")
			elif (i % 16) == 15:
				if (i == 511):
					OutputFile.write("{:>4s}".format(SpdTable[i]) + " //{:d}".format(i) + "\n")
				else:
					OutputFile.write("{:>4s}".format(SpdTable[i]) + "," + "//{:d}".format(i) + "\n")
			else:
				OutputFile.write("{:>4s}".format(SpdTable[i]) + ",")
		OutputFile.write("},\n")
		OutputFile.close()
		return 0

class Hynix:
	def __init__(self, ExcelFile):
		self._ExcelFile = ExcelFile
		self._SpdTable = []
	
	def __GetSpdTableFromFile(self):
		ExcelFile = xlrd.open_workbook(self._ExcelFile)
		ExcelSheet = ExcelFile.sheet_by_name("grd_excel")
		
		#initialize SpdTable
		SpdTable = []
		for i in range(512):
			SpdTable.append(0)
		
		#Fill in SpdTable
		for Index in range(0,ExcelSheet.ncols):
			if ExcelSheet.row_values(0)[Index] == "BYTE":
				IndexCol = Index
			if ExcelSheet.row_values(0)[Index] == "HEX":
				ValCol = Index
		for RowNo in range(1,ExcelSheet.nrows):
			HexValue = "0x" +  str(ExcelSheet.row_values(RowNo)[ValCol])[:2]
			SpdTable[eval(ExcelSheet.row_values(RowNo)[IndexCol])] = HexValue
		
		#CRC check
		ReturnVal = _Crc16Check(SpdTable)
		if not ReturnVal == 0:
			return ReturnVal
		
		self._SpdTable = SpdTable
		return 0
		
	def SaveSpdTableToFile(self, OutputFile):
		ReturnVal = self.__GetSpdTableFromFile()
		if not ReturnVal == 0:
			return ReturnVal
			
		#dump SpdTable
		SpdTable = self._SpdTable
		OutputFile = open(OutputFile, 'w')
		OutputFile.write("{//HYNIX\n")
		OutputFile.write("//  ")
		for i in range(16):
			OutputFile.write("{:^5d}".format(i))
		OutputFile.write("\n")
		for i in range(len(SpdTable)):
			if (i % 16) == 0:
				OutputFile.write("    ")
				OutputFile.write("{:>4s}".format(SpdTable[i]) + ",")
			elif (i % 16) == 15:
				if (i == 511):
					OutputFile.write("{:>4s}".format(SpdTable[i]) + " //{:d}".format(i) + "\n")
				else:
					OutputFile.write("{:>4s}".format(SpdTable[i]) + "," + "//{:d}".format(i) + "\n")
			else:
				OutputFile.write("{:>4s}".format(SpdTable[i]) + ",")
		OutputFile.write("},\n")
		OutputFile.close()
		return 0

class Samsung:
	def __init__(self, TxtFile):
		self._TxtFile = TxtFile
		self._SpdTable = []
	
	def __GetSpdTableFromFile(self):
		with open(self._TxtFile,"r",encoding="utf-8") as FileData:
			Data = FileData.readlines()
			
		#Initialize SpdTable
		SpdTable = []
		for i in range(512):
			SpdTable.append(0)
			
		#Delete invalid lines
		ValidData = []
		for Index in range(0,len(Data)):
			LineSplit = Data[Index].split()
			if LineSplit == []:
				continue
			else:
				FirstPara = LineSplit[0].split("~")
				if FirstPara[0].isnumeric():
					ValidData.append(LineSplit)
		
		#Fill in SpdTable
		for Index in range(0,len(ValidData)):
			FirstPara = ValidData[Index][0]
			if FirstPara.isnumeric():
				ByteIndex = eval(FirstPara)
				HexValue = ValidData[Index][-1]
				SpdTable[ByteIndex] = "0x" + HexValue[:-1]
			else:
				ByteIndex = FirstPara.split("~")
				ByteIndexL = eval(ByteIndex[0])
				ByteIndexR = eval(ByteIndex[1])
				for i in range(ByteIndexL,ByteIndexR + 1):
					HexValue = ValidData[Index][-1]
					SpdTable[i] = "0x" + HexValue[:-1]
		
		#CRC check
		ReturnVal = _Crc16Check(SpdTable)
		if not ReturnVal == 0:
			return ReturnVal
			
		self._SpdTable = SpdTable
		return 0
		
	def SaveSpdTableToFile(self, OutputFile):
		ReturnVal = self.__GetSpdTableFromFile()
		if not ReturnVal == 0:
			return ReturnVal
			
		#dump SpdTable
		SpdTable = self._SpdTable
		OutputFile = open(OutputFile, 'w')
		OutputFile.write("{//SANSUNG\n")
		OutputFile.write("//  ")
		for i in range(16):
			OutputFile.write("{:^5d}".format(i))
		OutputFile.write("\n")
		for i in range(len(SpdTable)):
			if (i % 16) == 0:
				OutputFile.write("    ")
				OutputFile.write("{:>4s}".format(SpdTable[i]) + ",")
			elif (i % 16) == 15:
				if (i == 511):
					OutputFile.write("{:>4s}".format(SpdTable[i]) + " //{:d}".format(i) + "\n")
				else:
					OutputFile.write("{:>4s}".format(SpdTable[i]) + "," + "//{:d}".format(i) + "\n")
			else:
				OutputFile.write("{:>4s}".format(SpdTable[i]) + ",")
		OutputFile.write("},\n")
		OutputFile.close()
		return 0
		
def Usage():
	print("################################")
	print("# AConvGPIO Version: 1.00      #")
	print("# Author: Yuan_Deng@asus.com   #")
	print("################################")
	print("Example: AConvSPD.exe -i DDR4.docx [-o [SpdTable.h]]")
	print("Input_file Format: DOC, DOCX, RTF, XLSX, XLS, TXT")
	
def Main(argv=None):
	if argv is None:
		argv = sys.argv

	argc = len(sys.argv)
	if argc == 1:
		Usage()
		return 1
	
	InputFile = ""
	OutputFile = ""
	InputFileFormat = ""
	try:
		opts, args = getopt.getopt(sys.argv[1:], "ho:i:", ["help"])
		if opts == []:
			Usage()
			return 1
		else:
			for op, val in opts:
				if op in ("-h", "--help"):
					Usage()
					return 1
				elif op in ("-i"):
					InputFile = val
				elif op in ("-o"):
					OutputFile = val
	except:
		print("ERROR:Invalid parameter!")
		return 1
	
	InputFileFormat = InputFile.split('.')[-1].upper()
	if InputFileFormat == "DOCX" or InputFileFormat == "DOC" or InputFileFormat == "RTF":
		FileParse = Micron(InputFile)
	elif InputFileFormat == "XLSX" or InputFileFormat == "XLS":
		FileParse = Hynix(InputFile)
	elif InputFileFormat == "TXT":
		FileParse = Samsung(InputFile)
	else:
		print("ERROR:Unsupported input file format!")
		print("Tips: Input_file extensions must be displayed!")
		return 1
		
	if OutputFile == "":
		OutputFile = "SpdTable.h"
		
	ReturnValue = FileParse.SaveSpdTableToFile(OutputFile)
	print("SPD table generated!")
	return 0
	
if __name__ == '__main__':
	sys.exit(Main())	